import os
from datetime import datetime


def export_txt(text: str, filename: str, output_dir: str = "outputs") -> str:
    """Export transcription as plain text."""
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{filename}.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(f"Transcription Export\n")
        f.write(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}\n")
        f.write("=" * 50 + "\n\n")
        f.write(text)
    return out_path


def export_pdf(text: str, filename: str, metadata: dict, output_dir: str = "outputs") -> str:
    """Export transcription as PDF."""
    from fpdf import FPDF
    os.makedirs(output_dir, exist_ok=True)
    
    pdf = FPDF()
    pdf.add_page()
    
    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 12, "WhisperTask - Transcription", new_x="LMARGIN", new_y="NEXT")
    
    # Metadata
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"File: {metadata.get('original_filename', 'N/A')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Confidence: {metadata.get('confidence_score', 'N/A')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Processing Time: {metadata.get('processing_time', 'N/A')}s", new_x="LMARGIN", new_y="NEXT")
    
    pdf.ln(5)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    
    # Transcription text
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(0, 7, text)
    
    out_path = os.path.join(output_dir, f"{filename}.pdf")
    pdf.output(out_path)
    return out_path


def export_docx(text: str, filename: str, metadata: dict, output_dir: str = "outputs") -> str:
    """Export transcription as DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    
    # Title
    title = doc.add_heading("WhisperTask - Transcription", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Metadata table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Light List"
    
    meta_items = [
        ("File", metadata.get("original_filename", "N/A")),
        ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("Confidence Score", str(metadata.get("confidence_score", "N/A"))),
        ("Processing Time", f"{metadata.get('processing_time', 'N/A')}s"),
    ]
    
    for i, (key, val) in enumerate(meta_items):
        row = meta_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
    
    doc.add_paragraph()
    doc.add_heading("Transcription", level=2)
    
    # Main text
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(12)
    
    out_path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(out_path)
    return out_path
